from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class DocxRenderer:
    def __init__(self):
        self.document = None

    def add_title(self, title: str) -> None:
        """Add title to the document."""
        if not isinstance(title, str):
            title = str(title)
        heading = self.document.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_query(self, query: str) -> None:
        """Add query section."""
        if isinstance(query, (tuple, list)):
            query = " ".join(str(q) for q in query)
        elif not isinstance(query, str):
            query = str(query)
        self.document.add_heading('Query', level=2)
        paragraph = self.document.add_paragraph(query)
        paragraph.style = 'Quote'

    def add_context(self, context: List[Dict[str, Any]]) -> None:
        """Add context section."""
        self.document.add_heading('Context', level=2)
        for doc in context:
            # Add source information
            source = doc.get('metadata', {}).get('source', 'Unknown')
            if not isinstance(source, str):
                source = str(source)
            source_para = self.document.add_paragraph()
            source_para.add_run(f"Source: {source}").italic = True

            # Add content
            text = doc.get('text', '')
            if isinstance(text, (tuple, list)):
                text = " ".join(str(t) for t in text)
            elif not isinstance(text, str):
                text = str(text)
            self.document.add_paragraph(text)
            self.document.add_paragraph()  # Add spacing

    def add_response(self, response: str) -> None:
        """Add response section."""
        self.document.add_heading('Response', level=2)
        # Ensure response is a string
        if isinstance(response, (tuple, list)):
            response = " ".join(str(r) for r in response)
        elif not isinstance(response, str):
            response = str(response)
        self.document.add_paragraph(response)

    def add_metadata(self, metadata: Dict[str, Any]) -> None:
        """Add metadata section."""
        self.document.add_heading('Metadata', level=2)
        for key, value in metadata.items():
            self.document.add_paragraph(f"{key}: {value}")

    def render(
        self,
        query: str,
        context: List[Dict[str, Any]],
        response: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """Render complete document."""
        self.document = Document()  # Ensure a new document each time

        # Add timestamped title
        self.add_title(f"Document Analysis - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Add main content
        self.add_query(query)
        self.add_context(context)
        self.add_response(response)

        # Add metadata if provided
        if metadata:
            self.add_metadata(metadata)

        return self.document

    def save(self, filename: str) -> None:
        """Save document to file."""
        if self.document is not None:
            self.document.save(filename)